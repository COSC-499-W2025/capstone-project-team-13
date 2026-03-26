"""
Thin export layer — PDF and DOCX rendering only.
Data is gathered via existing modules:
  - src.Portfolio.portfolioFormatter : PortfolioFormatter._infer_skills()
  - src.Databases.database           : db_manager
Deps: pip install reportlab python-docx
"""
from __future__ import annotations
import io, re
from typing import Any

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.Databases.database import db_manager
from src.Portfolio.portfolioFormatter import PortfolioFormatter
from src.UserPrompts.config_integration import has_ai_consent

_UUID_RE = re.compile(r"^[0-9a-f\-]{36}$", re.IGNORECASE)
_GREEN   = "#2d6a4f"


# ── helpers ───────────────────────────────────────────────────────────────────

def _skill_level(count: int) -> str:
    if count >= 5: return "Expert"
    if count >= 2: return "Proficient"
    return "Familiar"

def _clean_header(raw: str, fallback: str) -> str:
    if " | " in raw:
        parts = raw.split(" | ", 1)
        if _UUID_RE.match(parts[0].strip()):
            return parts[1]
    return raw or fallback

def _role_line(education: list) -> str:
    if not education:
        return ""
    e = education[0]
    return f"{e.get('degree_type','')} in {e.get('topic','')}".strip(" in")


def _try_generate_ai_bullets(project) -> list[str]:
    """
    Attempt to generate AI resume bullets for a project and save them.
    Returns the generated bullets, or [] if AI is unavailable or fails.
    Called only when has_ai_consent() is True and the project has no stored bullets.
    """
    try:
        from src.AI.ai_enhanced_summarizer import generate_resume_bullets
        from src.Resume.resumeAnalytics import score_all_bullets

        # Build a project dict matching what generate_resume_bullets expects
        project_dict = {
            "project_name": project.name,
            "skills": [s.strip() for s in (project.skills or "").split(",") if s.strip()]
                      if isinstance(project.skills, str)
                      else (project.skills or []),
            "file_count": project.file_count,
            "lines_of_code": project.lines_of_code,
            "success_score": getattr(project, "importance_score", 0) or 0,
            "contribution_score": getattr(project, "contribution_score", 0) or 0,
        }

        bullets = generate_resume_bullets(project_dict, num_bullets=3)
        if not bullets:
            return []

        # Persist so subsequent loads don't regenerate
        scoring = score_all_bullets(bullets, project.project_type)
        db_manager.save_resume_bullets(
            project_id=project.id,
            bullets=bullets,
            header=project.name,
            ats_score=scoring["overall_score"],
        )
        return bullets

    except Exception:
        return []


# ── data gathering ────────────────────────────────────────────────────────────

def get_resume_preview_data(user_id: int) -> dict[str, Any]:
    """Assemble resume data using existing helpers. Returns plain dict."""
    user = db_manager.get_user(user_id)
    if not user:
        raise ValueError(f"User {user_id} not found")

    education = db_manager.get_education_for_user(user_id)
    awards: list[str] = (user.resume or {}).get("awards", [])

    formatter    = PortfolioFormatter()
    all_projects = db_manager.get_all_projects(user_id=user_id)

    skill_counts: dict[str, int] = {}
    for p in all_projects:
        for skill in formatter._infer_skills(p):
            skill_counts[skill] = skill_counts.get(skill, 0) + 1

    skills_by_level: dict[str, list[str]] = {"Expert": [], "Proficient": [], "Familiar": []}
    for skill, count in sorted(skill_counts.items()):
        skills_by_level[_skill_level(count)].append(skill)

    ai_enabled = has_ai_consent()
    has_bullets = {p.id for p in all_projects if p.bullets is not None}
    resume_projects = []

    for project in all_projects:
        if project.id in has_bullets:
            bd     = db_manager.get_resume_bullets(project.id) or {}
            header = _clean_header(bd.get("header", ""), project.name)
            bullets = bd.get("bullets", [])
            ats     = bd.get("ats_score")
        else:
            header = project.name
            ats    = None
            # Auto-generate AI bullets if consent is granted
            if ai_enabled:
                bullets = _try_generate_ai_bullets(project)
            else:
                bullets = []

        resume_projects.append({
            "name":      project.name,
            "header":    header,
            "bullets":   bullets,
            "ats_score": ats,
        })

    return {
        "name":           f"{user.first_name} {user.last_name}",
        "email":          user.email,
        "education":      [e.to_dict() for e in education],
        "awards":         awards,
        "skills_by_level": skills_by_level,
        "projects":       resume_projects,
    }


# ── PDF builder ───────────────────────────────────────────────────────────────

def _build_pdf(data: dict[str, Any]) -> bytes:
    buf = io.BytesIO()

    BLACK = colors.black

    name_style    = ParagraphStyle("name",    fontSize=18, textColor=BLACK,
                                   alignment=TA_CENTER, fontName="Times-Bold",
                                   spaceAfter=2, spaceBefore=0)
    contact_style = ParagraphStyle("contact", fontSize=10, textColor=BLACK,
                                   alignment=TA_CENTER, spaceAfter=6)
    section_style = ParagraphStyle("section", fontSize=10, textColor=BLACK,
                                   fontName="Times-Bold", spaceBefore=8, spaceAfter=1)
    body_style    = ParagraphStyle("body",    fontSize=10, leading=14, spaceAfter=2)
    label_style   = ParagraphStyle("label",   fontSize=10, fontName="Times-Bold",
                                   spaceAfter=1)
    italic_style  = ParagraphStyle("italic",  fontSize=10, fontName="Times-Italic",
                                   spaceAfter=2)

    def rule():
        return HRFlowable(width="100%", thickness=0.5, color=BLACK,
                          spaceAfter=2, spaceBefore=0)

    def section_block(title):
        """Bold uppercase title followed by a horizontal rule."""
        return [Paragraph(f"<b>{title.upper()}</b>", section_style), rule()]

    def two_col(left, right):
        row = [[Paragraph(f"<b>{left}</b>", label_style),
                Paragraph(right, ParagraphStyle("r", fontSize=10, alignment=TA_RIGHT))]]
        t = Table(row, colWidths=[4.5*inch, 2.5*inch])
        t.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "BOTTOM"),
                               ("BOTTOMPADDING", (0,0), (-1,-1), 0)]))
        return t

    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.6*inch, bottomMargin=0.6*inch)
    story = []

    # Name + contact
    story.append(Paragraph(data.get("name", ""), name_style))
    contact_parts = [p for p in [
        data.get("phone"), data.get("email"),
        data.get("linkedin"), data.get("github")
    ] if p]
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))
    story.append(rule())

    # Render sections in user-defined order, falling back to legacy order
    section_order = data.get("section_order") or ["education", "awards", "skills", "work_history", "projects"]
    section_labels = data.get("section_labels") or {}
    skills_by_level = data.get("skills_by_level", {})

    for key in section_order:
        if key == "education" and data.get("education"):
            label = section_labels.get("education", "Education")
            story += section_block(label)
            for edu in data.get("education", []):
                end   = (edu.get("end_date") or "")
                end   = end[:7].replace("-", "/") if end and end != "Present" else "Present"
                inst  = edu.get("institution", "")
                loc   = edu.get("location", "")
                story.append(two_col(f"{inst}  |  {loc}" if loc else inst, end))
                degree = f"{edu.get('degree_type','')} in {edu.get('topic','')}".strip(" in")
                gpa    = edu.get("gpa", "")
                story.append(Paragraph(f"<i>{degree}{('  |  GPA: ' + gpa) if gpa else ''}</i>", italic_style))
                for detail in edu.get("details", []):
                    story.append(Paragraph(f"• {detail}", body_style))
            story.append(Spacer(1, 4))

        elif key == "awards" and data.get("awards"):
            label = section_labels.get("awards", "Awards")
            story += section_block(label)
            for award in data.get("awards", []):
                story.append(Paragraph(f"• {award}", body_style))
            story.append(Spacer(1, 4))

        elif key == "skills" and (any(v for v in skills_by_level.values()) or data.get("skills")):
            label = section_labels.get("skills", "Technical Skills")
            story += section_block(label)
            for level, items in skills_by_level.items():
                if items:
                    story.append(Paragraph(f"<b>{level}:</b> {', '.join(items)}", body_style))
            for line in data.get("skills", []):
                story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 4))

        elif key == "work_history" and data.get("work_history"):
            label = section_labels.get("work_history", "Relevant Experience")
            story += section_block(label)
            for job in data.get("work_history", []):
                comp  = job.get("company", "")
                loc   = job.get("location", "")
                start = job.get("start_date", "")
                end   = job.get("end_date", "Present")
                story.append(two_col(f"{comp}  |  {loc}" if loc else comp,
                                     f"{start} - {end}" if start else end))
                if job.get("role"):
                    story.append(Paragraph(f"<i>{job['role']}</i>", italic_style))
                for b in job.get("bullets", []):
                    story.append(Paragraph(f"• {b}", body_style))
                story.append(Spacer(1, 4))

        elif key == "projects" and data.get("projects"):
            label = section_labels.get("projects", "Projects")
            story += section_block(label)
            for proj in data.get("projects", []):
                header = proj.get("name") or proj.get("header") or ""
                story.append(two_col(header, proj.get("date", "")))
                for b in proj.get("bullets", []):
                    story.append(Paragraph(f"• {b}", body_style))
                story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()


# ── DOCX builder ─────────────────────────────────────────────────────────────────────────────────

def _build_docx(data: dict[str, Any]) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Pt(43)
        section.bottom_margin = Pt(43)
        section.left_margin   = Pt(54)
        section.right_margin  = Pt(54)

    def add_rule(p):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_para(text="", bold=False, italic=False, size=10,
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, space_before=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        p.alignment = align
        if text:
            r = p.add_run(text)
            r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        return p

    def section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(title.upper())
        r.bold = True; r.font.size = Pt(10)
        add_rule(p)

    def two_col_row(left, right, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(2)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab  = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), "8640")
        tabs.append(tab); pPr.append(tabs)
        r1 = p.add_run(left);  r1.bold = True;  r1.font.size = Pt(size)
        p.add_run("\t")
        r2 = p.add_run(right); r2.font.size = Pt(size)
        return p

    def bullet(text, size=10):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text); r.font.size = Pt(size)

    # Name + contact
    add_para(data.get("name", ""), bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    contact_parts = [p for p in [
        data.get("phone"), data.get("email"),
        data.get("linkedin"), data.get("github")
    ] if p]
    if contact_parts:
        add_para(" | ".join(contact_parts), size=10,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p_rule = add_para(space_after=4)
    add_rule(p_rule)

    # Render sections in user-defined order, falling back to legacy order
    section_order = data.get("section_order") or ["education", "awards", "skills", "work_history", "projects"]
    section_labels = data.get("section_labels") or {}
    skills_by_level = data.get("skills_by_level", {})

    for key in section_order:
        if key == "education" and data.get("education"):
            label = section_labels.get("education", "Education")
            section_heading(label)
            for edu in data.get("education", []):
                end   = (edu.get("end_date") or "")
                end   = end[:7].replace("-", "/") if end and end != "Present" else "Present"
                inst  = edu.get("institution", "")
                loc   = edu.get("location", "")
                two_col_row(f"{inst}  |  {loc}" if loc else inst, end)
                degree = f"{edu.get('degree_type','')} in {edu.get('topic','')}".strip(" in")
                gpa    = edu.get("gpa", "")
                add_para(degree + (f"  |  GPA: {gpa}" if gpa else ""),
                         italic=True, size=10, space_after=1)
                for detail in edu.get("details", []):
                    bullet(detail)

        elif key == "awards" and data.get("awards"):
            label = section_labels.get("awards", "Awards")
            section_heading(label)
            for award in data.get("awards", []):
                bullet(award)

        elif key == "skills" and (any(v for v in skills_by_level.values()) or data.get("skills")):
            label = section_labels.get("skills", "Technical Skills")
            section_heading(label)
            for level, items in skills_by_level.items():
                if items:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(1)
                    r1 = p.add_run(f"{level}: "); r1.bold = True; r1.font.size = Pt(10)
                    p.add_run(", ".join(items)).font.size = Pt(10)
            for line in data.get("skills", []):
                add_para(line, size=10, space_after=1)

        elif key == "work_history" and data.get("work_history"):
            label = section_labels.get("work_history", "Relevant Experience")
            section_heading(label)
            for job in data.get("work_history", []):
                comp  = job.get("company", "")
                loc   = job.get("location", "")
                start = job.get("start_date", "")
                end   = job.get("end_date", "Present")
                two_col_row(f"{comp}  |  {loc}" if loc else comp,
                            f"{start} - {end}" if start else end)
                if job.get("role"):
                    add_para(job["role"], italic=True, size=10, space_after=1)
                for b in job.get("bullets", []):
                    bullet(b)

        elif key == "projects" and data.get("projects"):
            label = section_labels.get("projects", "Projects")
            section_heading(label)
            for proj in data.get("projects", []):
                header = proj.get("name") or proj.get("header") or ""
                two_col_row(header, proj.get("date", ""))
                for b in proj.get("bullets", []):
                    bullet(b)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── public API ────────────────────────────────────────────────────────────────

def generate_resume_pdf(user_id: int) -> bytes:
    """Export the stored user.resume as PDF. Raises ValueError if no resume generated yet."""
    user = db_manager.get_user(user_id)
    if not user or not user.resume:
        raise ValueError("No resume found. Call POST /resume/generate first.")
    return _build_pdf(user.resume)

def generate_resume_docx(user_id: int) -> bytes:
    """Export the stored user.resume as DOCX. Raises ValueError if no resume generated yet."""
    user = db_manager.get_user(user_id)
    if not user or not user.resume:
        raise ValueError("No resume found. Call POST /resume/generate first.")
    return _build_docx(user.resume)